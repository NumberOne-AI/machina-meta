#!/usr/bin/env -S uv run --quiet --script
# /// script
# requires-python = ">=3.13"
# dependencies = ["google-genai", "python-docx", "pdf2image", "pillow", "spire-doc-free", "plum-dispatch==1.7.4"]
# ///
"""
Gemini DOCX Processing Proof of Concept

Demonstrates extracting structured biomedical/genetic information from DOCX files
using Google Gemini API with multiple extraction modes.

MODES:
    text    - Extract text from DOCX and send as plain text (fast, no formatting)
    image   - Convert DOCX to images and use Gemini vision (preserves layout)
    compare - Run both modes and compare results

EXTRACTION-ONLY MODES (no LLM, saves intermediate files):
    extract-text   - Extract text only, save to .txt file
    extract-images - Convert to images only, save as PNG files
    extract-all    - Extract both text and images, save all files

Usage:
    ./scripts/gemini_docx_poc.py <path-to-docx-file>
    ./scripts/gemini_docx_poc.py <path-to-docx-file> --mode=image
    ./scripts/gemini_docx_poc.py <path-to-docx-file> --mode=compare
    ./scripts/gemini_docx_poc.py <path-to-docx-file> --model=gemini-2.5-pro-preview-05-06

Image mode conversion (tries in order):
    1. Spire.Doc (pure Python, no system deps) - included in script dependencies
    2. LibreOffice + Poppler (system deps required):
       - macOS: brew install libreoffice poppler
       - Linux: apt install libreoffice poppler-utils

References:
    - https://ai.google.dev/gemini-api/docs/document-processing
    - https://ai.google.dev/api/files
    - https://github.com/eiceblue/Spire.Doc-for-Python
"""

import argparse
import base64
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from io import BytesIO
from pathlib import Path

from docx import Document
from google import genai
from google.genai import types
from PIL import Image


def load_env_variable(var_name: str, env_file: str = ".env") -> str | None:
    """Load environment variable from env file or environment.

    Priority:
    1. Environment variable (if set)
    2. .env file in current directory
    3. .env file in workspace root
    """
    # Check environment variable first (highest priority)
    value = os.getenv(var_name)
    if value:
        return value

    # Try .env in current working directory and workspace root
    env_paths = [
        Path.cwd() / env_file,
        Path(__file__).parent.parent / env_file,  # Workspace root
    ]

    for env_path in env_paths:
        if env_path.exists():
            try:
                with open(env_path) as f:
                    for line in f:
                        line = line.strip()
                        # Skip comments and empty lines
                        if not line or line.startswith("#"):
                            continue
                        # Match: VAR_NAME=value
                        match = re.match(rf"^{re.escape(var_name)}=(.*)$", line)
                        if match:
                            return match.group(1).strip().strip('"').strip("'")
            except OSError:
                continue

    return None


# Structured extraction prompt for biomedical/genetic data
EXTRACTION_PROMPT = """
Analyze the following medical/genetic document and extract ALL structured information into JSON format.

Extract the following categories if present:

1. **patient_info**: name, date_of_birth, gender, patient_id, collection_date, report_date
2. **biomarkers**: array of objects with fields: name, value, unit, reference_range, status (normal/high/low/critical), category
3. **genetic_variants**: array of objects with fields: gene, variant, zygosity, classification (pathogenic/benign/VUS), associated_conditions
4. **diagnoses**: array of objects with fields: condition, icd_code, status, onset_date
5. **medications**: array of objects with fields: name, dosage, frequency, route, start_date
6. **lab_panels**: array of objects with fields: panel_name, tests (array of name, value, unit, reference_range, flag)
7. **risk_scores**: array of objects with fields: name, score, interpretation, percentile
8. **recommendations**: array of objects with fields: type, description, priority, follow_up_date

Rules:
- Return ONLY valid JSON, no markdown code blocks
- Use null for missing values, not empty strings
- Preserve exact numeric values and units as written
- Include ALL data found, even if categories overlap
- For ranges, use format "min-max" or "< value" or "> value"
"""

TEXT_MODE_PROMPT = EXTRACTION_PROMPT + """
DOCUMENT TEXT:
---
{document_text}
---

Return the structured JSON:
"""

IMAGE_MODE_PROMPT = EXTRACTION_PROMPT + """
The document pages are provided as images. Analyze ALL pages carefully to extract the information.

Return the structured JSON:
"""


def extract_text_from_docx(docx_path: str) -> str:
    """Extract all text content from a DOCX file."""
    doc = Document(docx_path)

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_text))
        if table_text:
            text_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]")

    return "\n".join(text_parts)


def convert_docx_to_images_spire(docx_path: str, dpi: int = 300) -> list[Image.Image] | None:
    """Convert DOCX to images using Spire.Doc (pure Python, no system dependencies).

    Note: Spire.Doc free version does not support custom DPI for SaveImageToStreams.
    Images are generated at ~96 DPI. For higher resolution, use LibreOffice method.

    Args:
        docx_path: Path to DOCX file
        dpi: Requested resolution (note: Spire.Doc ignores this, outputs ~96 DPI)

    Returns:
        List of PIL Image objects, one per page, or None if Spire.Doc unavailable
    """
    # Skip Spire.Doc if high DPI is requested - it doesn't support custom resolution
    if dpi > 150:
        print(f"[WARN] Spire.Doc doesn't support {dpi} DPI, skipping to LibreOffice...")
        return None

    try:
        from spire.doc import Document as SpireDocument, ImageType
        from io import BytesIO
    except ImportError:
        return None  # Spire.Doc not available, fall back to other methods

    try:
        doc = SpireDocument()
        doc.LoadFromFile(docx_path)

        # Convert all pages to images (fixed ~96 DPI, Spire.Doc doesn't support custom DPI)
        image_streams = doc.SaveImageToStreams(ImageType.Bitmap)

        images = []
        for stream in image_streams:
            # Convert Spire stream to PIL Image
            image_data = stream.ToArray()
            pil_image = Image.open(BytesIO(bytes(image_data)))
            images.append(pil_image.copy())  # Copy to detach from stream

        doc.Close()
        return images
    except Exception as e:
        print(f"[WARN] Spire.Doc conversion failed: {e}")
        return None


def convert_docx_to_pdf_libreoffice(docx_path: str, output_dir: str) -> str | None:
    """Convert DOCX to PDF using LibreOffice headless mode.

    Args:
        docx_path: Path to DOCX file
        output_dir: Directory for PDF output

    Returns:
        Path to generated PDF file, or None if LibreOffice unavailable
    """
    # Check if LibreOffice is available
    libreoffice_cmd = None
    for cmd in ["libreoffice", "soffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice"]:
        if shutil.which(cmd):
            libreoffice_cmd = cmd
            break

    if not libreoffice_cmd:
        return None  # LibreOffice not available

    # Convert to PDF
    result = subprocess.run(
        [
            libreoffice_cmd,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ],
        capture_output=True,
        text=True
    )

    if result.returncode != 0:
        print(f"[WARN] LibreOffice conversion failed: {result.stderr}")
        return None

    # Find the generated PDF
    docx_name = Path(docx_path).stem
    pdf_path = Path(output_dir) / f"{docx_name}.pdf"

    if not pdf_path.exists():
        return None

    return str(pdf_path)


def convert_pdf_to_images(pdf_path: str, dpi: int = 200) -> list[Image.Image]:
    """Convert PDF pages to PIL Images.

    Args:
        pdf_path: Path to PDF file
        dpi: Resolution for rendering (higher = better quality but larger)

    Returns:
        List of PIL Image objects, one per page
    """
    from pdf2image import convert_from_path

    images = convert_from_path(pdf_path, dpi=dpi)
    return images


def convert_docx_to_images(docx_path: str, dpi: int = 300) -> tuple[list[Image.Image], str]:
    """Convert DOCX to images using best available method.

    Tries methods in order:
    1. Spire.Doc (pure Python, no system deps)
    2. LibreOffice + pdf2image (system deps required)

    Args:
        docx_path: Path to DOCX file
        dpi: Resolution for rendering (default 300 - industry standard for OCR)

    Returns:
        Tuple of (list of PIL Images, method name used)

    Raises:
        RuntimeError: If no conversion method is available
    """
    # Method 1: Try Spire.Doc (pure Python)
    print(f"[IMAGE MODE] Trying Spire.Doc conversion at {dpi} DPI...")
    images = convert_docx_to_images_spire(docx_path, dpi)
    if images:
        return images, "spire-doc"

    # Method 2: Try LibreOffice + pdf2image
    print("[IMAGE MODE] Spire.Doc unavailable, trying LibreOffice...")
    with tempfile.TemporaryDirectory() as temp_dir:
        pdf_path = convert_docx_to_pdf_libreoffice(docx_path, temp_dir)
        if pdf_path:
            try:
                images = convert_pdf_to_images(pdf_path, dpi=dpi)
                return images, "libreoffice"
            except Exception as e:
                print(f"[WARN] pdf2image conversion failed: {e}")

    # No methods available
    raise RuntimeError(
        "No DOCX to image conversion method available.\n"
        "Install one of:\n"
        "  - Spire.Doc: pip install spire-doc-free (pure Python, recommended)\n"
        "  - LibreOffice + Poppler:\n"
        "    - macOS: brew install libreoffice poppler\n"
        "    - Linux: apt install libreoffice poppler-utils"
    )


def image_to_base64(image: Image.Image, format: str = "PNG") -> str:
    """Convert PIL Image to base64 string."""
    buffer = BytesIO()
    image.save(buffer, format=format)
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


def extract_with_gemini_text(document_text: str, api_key: str, model: str = "gemini-3-pro-preview") -> dict:
    """
    Send document text to Gemini for structured extraction (text mode).

    Args:
        document_text: Extracted text from DOCX
        api_key: Google API key
        model: Gemini model to use

    Returns:
        Extracted structured data as dict
    """
    client = genai.Client(api_key=api_key)

    prompt = TEXT_MODE_PROMPT.format(document_text=document_text)

    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.1,  # Low temperature for consistent extraction
            response_mime_type="application/json",  # Request JSON response
        )
    )

    return _parse_json_response(response.text)


def extract_with_gemini_vision(images: list[Image.Image], api_key: str, model: str = "gemini-3-pro-preview") -> dict:
    """
    Send document images to Gemini for structured extraction (vision mode).

    Args:
        images: List of PIL Images (document pages)
        api_key: Google API key
        model: Gemini model to use

    Returns:
        Extracted structured data as dict
    """
    client = genai.Client(api_key=api_key)

    # Build multimodal content: images + prompt with high resolution for OCR
    parts = []

    # Add each page as an image part with media_resolution_high for better text extraction
    for i, image in enumerate(images):
        image_bytes = BytesIO()
        image.save(image_bytes, format="PNG")
        image_bytes.seek(0)

        parts.append(types.Part.from_bytes(
            data=image_bytes.read(),
            mime_type="image/png",
            video_metadata=None,  # Not a video
        ))

    # Add the extraction prompt (note: text is a keyword-only argument)
    parts.append(types.Part.from_text(text=IMAGE_MODE_PROMPT))

    response = client.models.generate_content(
        model=model,
        contents=parts,
        config=types.GenerateContentConfig(
            temperature=0.1,
            response_mime_type="application/json",
            media_resolution="high",  # Use high resolution for better OCR/text extraction
        )
    )

    return _parse_json_response(response.text)


def _parse_json_response(text: str) -> dict:
    """Parse JSON from Gemini response, handling markdown wrapping."""
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        # Try to extract JSON from response if wrapped in markdown
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]
        try:
            return json.loads(text.strip())
        except json.JSONDecodeError:
            # Return raw response for debugging
            return {"_raw_response": text, "_parse_error": str(e)}


def process_docx_text_mode(docx_path: str, api_key: str, model: str = "gemini-3-pro-preview") -> dict:
    """
    Process DOCX file using text extraction mode.

    Args:
        docx_path: Path to DOCX file
        api_key: Google API key
        model: Gemini model to use

    Returns:
        Dict with extraction results and metadata (includes processing_time_seconds)
    """
    start_time = time.time()
    path = Path(docx_path)

    # Extract text from DOCX
    print(f"[TEXT MODE] Extracting text from: {path.name}")
    document_text = extract_text_from_docx(str(path))

    if not document_text.strip():
        raise ValueError("No text content found in document")

    print(f"[TEXT MODE] Extracted {len(document_text)} characters")
    print(f"[TEXT MODE] Sending to Gemini ({model})...")

    # Extract structured data using Gemini
    extracted_data = extract_with_gemini_text(document_text, api_key, model)

    processing_time = time.time() - start_time
    print(f"[TEXT MODE] Completed in {processing_time:.2f} seconds")

    return {
        "mode": "text",
        "source_file": path.name,
        "text_length": len(document_text),
        "model_used": model,
        "processing_time_seconds": round(processing_time, 2),
        "extracted_data": extracted_data
    }


def process_docx_image_mode(docx_path: str, api_key: str, model: str = "gemini-3-pro-preview", dpi: int = 200) -> dict:
    """
    Process DOCX file using image/vision mode.

    Converts DOCX → Images → Gemini Vision
    Uses Spire.Doc (pure Python) or LibreOffice + pdf2image as fallback.

    Args:
        docx_path: Path to DOCX file
        api_key: Google API key
        model: Gemini model to use
        dpi: Image resolution (default 200, used by LibreOffice method)

    Returns:
        Dict with extraction results and metadata (includes processing_time_seconds)
    """
    start_time = time.time()
    path = Path(docx_path)

    print(f"[IMAGE MODE] Converting: {path.name}")

    # Convert DOCX to images using best available method
    conversion_start = time.time()
    images, conversion_method = convert_docx_to_images(str(path), dpi=dpi)
    conversion_time = time.time() - conversion_start
    print(f"[IMAGE MODE] Converted using: {conversion_method} ({conversion_time:.2f}s)")
    print(f"[IMAGE MODE] Generated {len(images)} page image(s)")

    # Calculate total image size
    total_pixels = sum(img.width * img.height for img in images)

    print(f"[IMAGE MODE] Sending to Gemini vision ({model})...")

    # Extract using vision
    llm_start = time.time()
    extracted_data = extract_with_gemini_vision(images, api_key, model)
    llm_time = time.time() - llm_start

    processing_time = time.time() - start_time
    print(f"[IMAGE MODE] Completed in {processing_time:.2f} seconds (conversion: {conversion_time:.2f}s, LLM: {llm_time:.2f}s)")

    return {
        "mode": "image",
        "source_file": path.name,
        "page_count": len(images),
        "total_pixels": total_pixels,
        "dpi": dpi,
        "conversion_method": conversion_method,
        "model_used": model,
        "processing_time_seconds": round(processing_time, 2),
        "conversion_time_seconds": round(conversion_time, 2),
        "llm_time_seconds": round(llm_time, 2),
        "extracted_data": extracted_data
    }


def extract_text_only(docx_path: str, output_dir: Path | None = None) -> dict:
    """Extract text from DOCX and save to file (no LLM processing).

    Args:
        docx_path: Path to DOCX file
        output_dir: Optional output directory (default: same as input file)

    Returns:
        Dict with extraction metadata
    """
    path = Path(docx_path)
    if output_dir is None:
        output_dir = path.parent

    print(f"[EXTRACT-TEXT] Processing: {path.name}")

    document_text = extract_text_from_docx(str(path))

    if not document_text.strip():
        print(f"[EXTRACT-TEXT] WARNING: No text content found")
        return {"source_file": path.name, "text_length": 0, "output_file": None}

    # Save text file
    output_path = output_dir / f"{path.stem}.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(document_text)

    print(f"[EXTRACT-TEXT] Saved: {output_path} ({len(document_text)} chars)")

    return {
        "source_file": path.name,
        "text_length": len(document_text),
        "output_file": str(output_path)
    }


def extract_images_only(docx_path: str, output_dir: Path | None = None, dpi: int = 300) -> dict:
    """Convert DOCX to images and save as PNG files (no LLM processing).

    Args:
        docx_path: Path to DOCX file
        output_dir: Optional output directory (default: creates subdir next to input)
        dpi: Image resolution (used by LibreOffice method)

    Returns:
        Dict with extraction metadata
    """
    path = Path(docx_path)
    if output_dir is None:
        output_dir = path.parent / f"{path.stem}_images"

    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"[EXTRACT-IMAGES] Processing: {path.name}")

    images, conversion_method = convert_docx_to_images(str(path), dpi=dpi)

    print(f"[EXTRACT-IMAGES] Converted using: {conversion_method}")
    print(f"[EXTRACT-IMAGES] Generated {len(images)} page(s)")

    saved_files = []
    for i, img in enumerate(images):
        img_path = output_dir / f"page_{i + 1:03d}.png"
        img.save(img_path, "PNG")
        saved_files.append(str(img_path))
        print(f"[EXTRACT-IMAGES] Saved: {img_path}")

    return {
        "source_file": path.name,
        "page_count": len(images),
        "conversion_method": conversion_method,
        "output_dir": str(output_dir),
        "output_files": saved_files
    }


def extract_all(docx_path: str, output_dir: Path | None = None, dpi: int = 300) -> dict:
    """Extract both text and images from DOCX (no LLM processing).

    Args:
        docx_path: Path to DOCX file
        output_dir: Optional output directory (default: creates subdir next to input)
        dpi: Image resolution

    Returns:
        Dict with extraction metadata for both text and images
    """
    path = Path(docx_path)
    if output_dir is None:
        output_dir = path.parent / f"{path.stem}_extracted"

    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print(f"EXTRACTING: {path.name}")
    print("=" * 60)

    # Extract text
    text_result = extract_text_only(str(path), output_dir)

    # Extract images
    images_dir = output_dir / "images"
    images_result = extract_images_only(str(path), images_dir, dpi)

    return {
        "source_file": path.name,
        "output_dir": str(output_dir),
        "text": text_result,
        "images": images_result
    }


def process_directory(
    input_dir: str,
    output_dir: str | None = None,
    mode: str = "extract-all",
    dpi: int = 300
) -> list[dict]:
    """Process all DOCX files in a directory.

    Args:
        input_dir: Directory containing DOCX files
        output_dir: Output directory (default: input_dir/extracted)
        mode: Extraction mode (extract-text, extract-images, extract-all)
        dpi: Image resolution

    Returns:
        List of results for each file
    """
    input_path = Path(input_dir)
    if not input_path.is_dir():
        raise ValueError(f"Not a directory: {input_dir}")

    docx_files = list(input_path.rglob("*.docx")) + list(input_path.rglob("*.doc"))
    if not docx_files:
        raise ValueError(f"No DOCX files found in: {input_dir}")

    if output_dir is None:
        output_path = input_path / "extracted"
    else:
        output_path = Path(output_dir)

    output_path.mkdir(parents=True, exist_ok=True)

    print(f"Found {len(docx_files)} DOCX file(s)")
    print(f"Output directory: {output_path}")
    print("=" * 60)

    results = []
    for docx_file in sorted(docx_files):
        # Create output subdir for each file
        file_output_dir = output_path / docx_file.stem

        if mode == "extract-text":
            result = extract_text_only(str(docx_file), file_output_dir)
        elif mode == "extract-images":
            result = extract_images_only(str(docx_file), file_output_dir, dpi)
        else:  # extract-all
            result = extract_all(str(docx_file), file_output_dir, dpi)

        results.append(result)
        print()

    return results


def compare_results(text_result: dict, image_result: dict) -> dict:
    """
    Compare extraction results between text and image modes.

    Args:
        text_result: Result from text mode
        image_result: Result from image mode

    Returns:
        Comparison analysis
    """
    text_data = text_result.get("extracted_data", {})
    image_data = image_result.get("extracted_data", {})

    comparison = {
        "source_file": text_result.get("source_file"),
        "model_used": text_result.get("model_used"),
        "text_mode": {
            "text_length": text_result.get("text_length"),
            "processing_time_seconds": text_result.get("processing_time_seconds"),
        },
        "image_mode": {
            "page_count": image_result.get("page_count"),
            "total_pixels": image_result.get("total_pixels"),
            "dpi": image_result.get("dpi"),
            "conversion_method": image_result.get("conversion_method"),
            "processing_time_seconds": image_result.get("processing_time_seconds"),
            "conversion_time_seconds": image_result.get("conversion_time_seconds"),
            "llm_time_seconds": image_result.get("llm_time_seconds"),
        },
        "extraction_counts": {},
        "differences": []
    }

    # Compare counts for each category
    categories = ["biomarkers", "genetic_variants", "diagnoses", "medications", "lab_panels", "risk_scores", "recommendations"]

    for cat in categories:
        text_items = text_data.get(cat, [])
        image_items = image_data.get(cat, [])

        text_count = len(text_items) if isinstance(text_items, list) else 0
        image_count = len(image_items) if isinstance(image_items, list) else 0

        comparison["extraction_counts"][cat] = {
            "text_mode": text_count,
            "image_mode": image_count,
            "difference": image_count - text_count
        }

        if text_count != image_count:
            comparison["differences"].append(f"{cat}: text={text_count}, image={image_count}")

    # Summary
    total_text = sum(comparison["extraction_counts"][c]["text_mode"] for c in categories)
    total_image = sum(comparison["extraction_counts"][c]["image_mode"] for c in categories)

    text_time = text_result.get("processing_time_seconds", 0)
    image_time = image_result.get("processing_time_seconds", 0)

    comparison["summary"] = {
        "total_items_text": total_text,
        "total_items_image": total_image,
        "difference": total_image - total_text,
        "winner": "image" if total_image > total_text else ("text" if total_text > total_image else "tie"),
        "text_processing_time_seconds": text_time,
        "image_processing_time_seconds": image_time,
        "time_ratio": round(image_time / text_time, 2) if text_time > 0 else None,
    }

    return comparison


def process_docx_file(docx_path: str, api_key: str, model: str = "gemini-3-pro-preview", mode: str = "text") -> dict:
    """
    Main function to process a DOCX file and extract structured data.

    Args:
        docx_path: Path to DOCX file
        api_key: Google API key
        model: Gemini model to use
        mode: Extraction mode ("text", "image", or "compare")

    Returns:
        Dict with extraction results and metadata
    """
    path = Path(docx_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    if path.suffix.lower() not in [".docx", ".doc"]:
        raise ValueError(f"Expected DOCX file, got: {path.suffix}")

    if mode == "text":
        return process_docx_text_mode(str(path), api_key, model)
    elif mode == "image":
        return process_docx_image_mode(str(path), api_key, model)
    elif mode == "compare":
        print("=" * 60)
        print("RUNNING TEXT MODE")
        print("=" * 60)
        text_result = process_docx_text_mode(str(path), api_key, model)

        print("\n" + "=" * 60)
        print("RUNNING IMAGE MODE")
        print("=" * 60)
        image_result = process_docx_image_mode(str(path), api_key, model)

        print("\n" + "=" * 60)
        print("COMPARING RESULTS")
        print("=" * 60)
        comparison = compare_results(text_result, image_result)

        return {
            "text_result": text_result,
            "image_result": image_result,
            "comparison": comparison
        }
    else:
        raise ValueError(f"Unknown mode: {mode}. Use 'text', 'image', or 'compare'")


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Extract biomedical/genetic data from DOCX files using Gemini",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
LLM Modes (require GOOGLE_API_KEY):
  text    - Extract text from DOCX and send as plain text (fast, no formatting)
  image   - Convert DOCX to images and use Gemini vision (preserves layout)
  compare - Run both modes and compare results

Extraction-Only Modes (no LLM, saves intermediate files):
  extract-text   - Extract text only, save to .txt file
  extract-images - Convert to images only, save as PNG files
  extract-all    - Extract both text and images, save all files

Image mode conversion (tries in order):
  1. Spire.Doc (pure Python, no system deps) - included in script dependencies
  2. LibreOffice + Poppler (system deps, fallback):
     - macOS: brew install libreoffice poppler
     - Linux: apt install libreoffice poppler-utils

Examples:
  # LLM extraction
  ./scripts/gemini_docx_poc.py document.docx
  ./scripts/gemini_docx_poc.py document.docx --mode=image
  ./scripts/gemini_docx_poc.py document.docx --mode=compare

  # Extraction only (no LLM)
  ./scripts/gemini_docx_poc.py document.docx --mode=extract-images
  ./scripts/gemini_docx_poc.py documents/ --mode=extract-all --output-dir=output/
        """
    )

    parser.add_argument("docx_path", help="Path to DOCX file or directory")
    parser.add_argument("--mode", "-m",
                        choices=["text", "image", "compare", "extract-text", "extract-images", "extract-all"],
                        default="text",
                        help="Extraction mode (default: text)")
    parser.add_argument("--model", default="gemini-3-pro-preview",
                        help="Gemini model to use (default: gemini-3-pro-preview)")
    parser.add_argument("--dpi", type=int, default=300,
                        help="Image resolution for image mode (default: 300 - industry standard for OCR)")
    parser.add_argument("--output-dir", "-o", default=None,
                        help="Output directory for extraction modes (default: next to input)")

    args = parser.parse_args()

    input_path = Path(args.docx_path)
    output_dir = Path(args.output_dir) if args.output_dir else None

    # Handle extraction-only modes (no LLM required)
    if args.mode.startswith("extract-"):
        try:
            if input_path.is_dir():
                # Process directory
                results = process_directory(
                    str(input_path),
                    str(output_dir) if output_dir else None,
                    args.mode,
                    args.dpi
                )
                print("=" * 60)
                print(f"COMPLETED: {len(results)} file(s) processed")
                print("=" * 60)
                print(json.dumps(results, indent=2))
            else:
                # Process single file
                if args.mode == "extract-text":
                    result = extract_text_only(str(input_path), output_dir)
                elif args.mode == "extract-images":
                    result = extract_images_only(str(input_path), output_dir, args.dpi)
                else:  # extract-all
                    result = extract_all(str(input_path), output_dir, args.dpi)

                print("\n" + "=" * 60)
                print("EXTRACTION RESULT:")
                print("=" * 60)
                print(json.dumps(result, indent=2))

        except Exception as e:
            import traceback
            print(f"ERROR: {e}")
            traceback.print_exc()
            sys.exit(1)

        return

    # LLM modes require API key
    api_key = load_env_variable("GOOGLE_API_KEY") or load_env_variable("GEMINI_API_KEY")
    if not api_key:
        print("ERROR: GOOGLE_API_KEY or GEMINI_API_KEY not found in environment or .env")
        print("Note: Use --mode=extract-images or --mode=extract-all for extraction without LLM")
        sys.exit(1)

    try:
        result = process_docx_file(args.docx_path, api_key, args.model, args.mode)

        print("\n" + "=" * 60)
        print("EXTRACTION RESULT:")
        print("=" * 60)
        print(json.dumps(result, indent=2))

        # Save to file with mode suffix
        docx_path = Path(args.docx_path)
        if args.mode == "compare":
            output_path = docx_path.with_suffix(".compare.json")
        else:
            output_path = docx_path.with_suffix(f".{args.mode}.json")

        with open(output_path, "w") as f:
            json.dump(result, f, indent=2)
        print(f"\nSaved to: {output_path}")

    except Exception as e:
        import traceback
        print(f"ERROR: {e}")
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
